import io

from docx import Document

from rios.core.schemas import (
    JournalRecommendation,
    MethodologyRecommendation,
    Paper,
    ResearchGapCandidate,
    SearchStrategy,
)
from rios.reports.docx_report import build_gap_report_docx


def _gap() -> ResearchGapCandidate:
    return ResearchGapCandidate(
        gap_id="g1", domain="Ag Econ", gap_type="methodological",
        description="A very specific candidate research gap description.",
        why_insufficient="Existing studies rely on a single method.",
        expected_contribution="A combined framework improving accuracy.",
        supporting_paper_ids=["W1"], confidence_score=0.75,
        prompt_version="v1", model_version="gemini-2.5-flash",
    )


def _paper() -> Paper:
    return Paper(
        id="W1", title="Forecasting Commodity Prices", authors=["A. Researcher"],
        year=2023, journal="Journal of Ag Econ", source="test",
    )


def _all_text(doc: Document) -> str:
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def test_report_contains_gap_description():
    gap = _gap()
    docx_bytes = build_gap_report_docx(gap, {"W1": _paper()}, [], [], None)
    doc = Document(io.BytesIO(docx_bytes))
    assert gap.description in _all_text(doc)


def test_report_contains_supporting_paper_title():
    gap = _gap()
    paper = _paper()
    docx_bytes = build_gap_report_docx(gap, {"W1": paper}, [], [], None)
    doc = Document(io.BytesIO(docx_bytes))
    assert paper.title in _all_text(doc)


def test_report_includes_methodology_section_when_present():
    gap = _gap()
    method = MethodologyRecommendation(
        name="Panel Data Analysis", mention_count=2, supporting_paper_ids=["W1"],
        typical_applications="app", strengths="strength", limitations="limit",
        reason_for_recommendation="Mentioned in 2 passages.",
    )
    docx_bytes = build_gap_report_docx(gap, {"W1": _paper()}, [method], [], None)
    doc = Document(io.BytesIO(docx_bytes))
    text = _all_text(doc)
    assert "Panel Data Analysis" in text
    assert "Mentioned in 2 passages." in text


def test_report_includes_journal_section_when_present():
    gap = _gap()
    journal = JournalRecommendation(
        journal_name="Journal of Ag Econ", paper_count=1, total_citations=5,
        supporting_paper_ids=["W1"], reason_for_recommendation="Published here before.",
    )
    docx_bytes = build_gap_report_docx(gap, {"W1": _paper()}, [], [journal], None)
    doc = Document(io.BytesIO(docx_bytes))
    assert "Journal of Ag Econ" in _all_text(doc)


def test_report_includes_search_strategy_appendix():
    gap = _gap()
    strategy = SearchStrategy(
        domain="Ag Econ", keywords=["commodity", "prices"],
        databases_searched=["OpenAlex"], publication_year_min=2015,
        publication_year_max=2026,
        journal_quartiles_allowed=[],
    )
    docx_bytes = build_gap_report_docx(gap, {"W1": _paper()}, [], [], strategy)
    doc = Document(io.BytesIO(docx_bytes))
    text = _all_text(doc)
    assert "OpenAlex" in text
    assert "commodity" in text


def test_report_handles_missing_papers_gracefully():
    gap = _gap()  # supporting_paper_ids=["W1"] but papers_by_id is empty
    docx_bytes = build_gap_report_docx(gap, {}, [], [], None)
    doc = Document(io.BytesIO(docx_bytes))
    assert "No supporting paper metadata available." in _all_text(doc)
