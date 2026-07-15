"""
Word (.docx) report generation for an accepted research gap.

WHY this module exists: everything RIOS produces before this point lives
inside the app session — useful for review, but not something a researcher
can hand to a supervisor or keep for their records. This module assembles
one gap's full trail (the gap itself, its supporting evidence, methodology
and journal recommendations, and the original search strategy) into a
single downloadable document.

WHY python-docx, generating bytes in memory (not writing to disk): the app
runs on Streamlit Cloud, where the filesystem is ephemeral and shouldn't be
relied on for anything the user needs to keep — building the document in an
io.BytesIO buffer and handing it straight to Streamlit's download button
is the correct pattern for a hosted app.

WHY nothing here is invented: every section either restates fields already
validated on the Pydantic models (ResearchGapCandidate, Paper,
MethodologyRecommendation, JournalRecommendation, SearchStrategy) or is
plain formatting — this module has zero decision-making logic of its own.
"""

from __future__ import annotations

import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from rios.core.schemas import (
    JournalRecommendation,
    MethodologyRecommendation,
    Paper,
    ResearchGapCandidate,
    SearchStrategy,
)


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _add_labeled_paragraph(doc: Document, label: str, value: str) -> None:
    p = doc.add_paragraph()
    run_label = p.add_run(f"{label}: ")
    run_label.bold = True
    p.add_run(value)


def build_gap_report_docx(
    gap: ResearchGapCandidate,
    papers_by_id: dict[str, Paper],
    methodology_recs: list[MethodologyRecommendation],
    journal_recs: list[JournalRecommendation],
    strategy: SearchStrategy | None,
) -> bytes:
    """Build a Word report for one accepted (or modified) gap. Returns raw
    .docx bytes, ready to hand to a download button — nothing is written
    to disk."""
    doc = Document()

    # --- Title & status ---
    title = doc.add_heading("RIOS Research Gap Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Research Intelligence Operating System — Built by Suman_Econ (UAS-B)")
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(10)

    doc.add_paragraph()  # spacer
    _add_labeled_paragraph(doc, "Domain", gap.domain)
    _add_labeled_paragraph(doc, "Gap type", gap.gap_type)
    _add_labeled_paragraph(doc, "Review status", gap.review_status.value.upper())
    if gap.review_comment:
        _add_labeled_paragraph(doc, "Reviewer comment", gap.review_comment)

    # --- Core gap content ---
    _add_heading(doc, "Research Gap", level=1)
    doc.add_paragraph(gap.description)

    _add_heading(doc, "Why Existing Evidence Is Insufficient", level=2)
    doc.add_paragraph(gap.why_insufficient)

    _add_heading(doc, "Expected Contribution", level=2)
    doc.add_paragraph(gap.expected_contribution)

    # --- Supporting evidence table ---
    _add_heading(doc, "Supporting Evidence", level=1)
    supporting_papers = [
        papers_by_id[pid] for pid in gap.supporting_paper_ids if pid in papers_by_id
    ]
    if supporting_papers:
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        header_cells = table.rows[0].cells
        for cell, text in zip(header_cells, ["Title", "Authors", "Year", "Journal"]):
            cell.text = text
        for paper in supporting_papers:
            row = table.add_row().cells
            row[0].text = paper.title
            row[1].text = ", ".join(paper.authors[:3]) or "Unknown"
            row[2].text = str(paper.year or "n.d.")
            row[3].text = paper.journal or "Unknown"
    else:
        doc.add_paragraph("No supporting paper metadata available.")

    # --- Methodology recommendations ---
    if methodology_recs:
        _add_heading(doc, "Methodology Recommendations", level=1)
        for m in methodology_recs:
            _add_heading(doc, m.name, level=2)
            _add_labeled_paragraph(doc, "Basis", m.reason_for_recommendation)
            _add_labeled_paragraph(doc, "Typical applications", m.typical_applications)
            _add_labeled_paragraph(doc, "Strengths", m.strengths)
            _add_labeled_paragraph(doc, "Limitations", m.limitations)
            if m.alternatives_considered:
                _add_labeled_paragraph(
                    doc, "Alternatives also seen in evidence",
                    ", ".join(m.alternatives_considered),
                )

    # --- Journal recommendations ---
    if journal_recs:
        _add_heading(doc, "Journal Recommendations", level=1)
        for j in journal_recs:
            _add_heading(doc, j.journal_name, level=2)
            _add_labeled_paragraph(
                doc, "Basis",
                f"{j.paper_count} supporting paper(s), {j.total_citations} total citations",
            )
            doc.add_paragraph(j.reason_for_recommendation)

    # --- Reproducibility appendix ---
    _add_heading(doc, "Appendix: Reproducibility Record", level=1)
    _add_labeled_paragraph(doc, "Prompt version", gap.prompt_version)
    _add_labeled_paragraph(doc, "Model version", gap.model_version)
    _add_labeled_paragraph(doc, "Generated at (UTC)", gap.generated_at.isoformat())
    _add_labeled_paragraph(doc, "Confidence score", f"{gap.confidence_score:.2f}")
    if strategy:
        _add_labeled_paragraph(doc, "Databases searched", ", ".join(strategy.databases_searched))
        _add_labeled_paragraph(doc, "Search keywords", ", ".join(strategy.keywords))
        _add_labeled_paragraph(
            doc, "Publication year range",
            f"{strategy.publication_year_min}–{strategy.publication_year_max}",
        )
        if strategy.inclusion_criteria:
            _add_labeled_paragraph(doc, "Inclusion criteria", "; ".join(strategy.inclusion_criteria))

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
